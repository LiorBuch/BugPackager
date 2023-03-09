import datetime
import json
import os
import sys
from tkinter.filedialog import askopenfile

from docx import Document

prod = True
nav_bar_height = 60
if not prod:
    exe_loc = os.path.dirname(sys.executable)
else:
    exe_loc = os.path.dirname(os.path.realpath(__file__))


def write_to_json(file_to_modify, data_list: list):
    with open(os.path.join(exe_loc, file_to_modify), "r") as jsonFile:
        data = json.load(jsonFile)
        for item in data_list:
            data[item[0]][item[1]] = item[2]
    with open(os.path.join(exe_loc, file_to_modify), "w") as jsonFile:
        json.dump(data, jsonFile)


def read_from_json(file, *args):
    with open(os.path.join(exe_loc, file), "r") as jsonFile:
        data = json.load(jsonFile)
        for index in args:
            data = data[index]
        return data


def create_doc(title, body, company, username):
    doc = Document()
    doc.add_heading(f"Bug Report From:{company}")
    doc.add_paragraph(f"A bug report made by:{username} at:{datetime.date.today()}")
    doc.add_heading(title, level=1)
    doc.add_paragraph(body)
    # doc.add_picture('monty-truth.png')
    doc.add_page_break()
    return doc
    doc.save("output\\temp.docx")
